#!/usr/bin/env python3
# VARIANT-GNN — Missense Varyant Patojenite Tahmini
# Telif Hakkı (c) 2026 XYRA3 Takımı. Tüm hakları saklıdır.
# TEKNOFEST 2026 Sağlıkta Yapay Zeka Yarışması (Üniversite ve Üzeri).
# Lisans: depo kökündeki LICENSE. Yalnızca araştırma/eğitim/yarışma amaçlıdır;
# klinik tanı/tedavi için kullanılamaz (Şartname §10).

"""
scripts/build_pdr_docx.py
=========================
reports/PDR_VARIANT_GNN_2026.md  ->  reports/PDR_VARIANT_GNN_2026.docx

TASLAK donusturucu. python-docx ile baslik/paragraf/tablo/madde/figur-gomme +
TEKNOFEST format hedefi (Aptos 12pt govde / 14pt baslik, 1.15 satir, justified,
ust 2.8 cm / diger 2.5 cm marj).

UYARI: Bu bir TASLAK uretir. Resmi sablona birebir uyum, <=10 sayfa siniri ve
gorsel yerlesim Word'de MANUEL dogrulanmalidir. Tablolar/figurler sayfa tasirsa
elle ayarlanir. Aptos sistemde yoksa Word ikame font kullanir.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "reports" / "PDR_VARIANT_GNN_2026.md"
OUT = ROOT / "reports" / "PDR_VARIANT_GNN_2026.docx"
FIG_DIR = ROOT

FONT = "Aptos"
SEP_RE = re.compile(r"^\s*:?-{2,}:?\s*$")
PNG_RE = re.compile(r"(reports/figures/[\w/]+\.png)")


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    pf = st.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(3)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for sec in doc.sections:
        sec.top_margin = Cm(2.8)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)


def _add_runs(par, text):
    """**bold** ve `code` isaretlerini Word run'larina cevir (basit)."""
    text = text.replace("`", "")
    for chunk in re.split(r"(\*\*.+?\*\*)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = par.add_run(chunk[2:-2])
            run.bold = True
        else:
            par.add_run(chunk)


def _heading(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level <= 1 else 6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.strip())
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt({0: 16, 1: 14, 2: 13, 3: 12}.get(level, 12))
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def _table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            par = cells[ci].paragraphs[0]
            par.text = ""
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(par, txt.strip())
            for run in par.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True


def _figure(doc, path_str, caption):
    img = FIG_DIR / path_str
    if not img.exists():
        return False
    try:
        doc.add_picture(str(img), width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption.strip())
            run.italic = True
            run.font.size = Pt(9)
        return True
    except Exception as exc:  # noqa: BLE001
        print(f"  figur gomulemedi ({path_str}): {exc}")
        return False


def _split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")  # Windows cp1254 stdout -> utf-8
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    _set_base_style(doc)

    i, n = 0, len(lines)
    n_tab = n_fig = 0
    while i < n:
        stripped = lines[i].strip()

        # Tablo blogu
        if stripped.startswith("|") and "|" in stripped[1:]:
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = []
            for bl in block:
                cells = _split_row(bl)
                if all(SEP_RE.match(c) or c == "" for c in cells):
                    continue  # ayrac satiri
                rows.append(cells)
            _table(doc, rows)
            n_tab += 1
            continue

        # Baslik
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            _heading(doc, m.group(2), len(m.group(1)) - 1)
            i += 1
            continue

        # Sayfa sonu (kapak/icindekiler muafiyeti icin)
        if stripped == "[[PAGEBREAK]]":
            doc.add_page_break()
            i += 1
            continue

        # Yatay cizgi
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Figur (png yolu iceren satir)
        png = PNG_RE.search(stripped)
        if png:
            caption = re.sub(r"\*|`", "", stripped)
            if _figure(doc, png.group(1), caption):
                n_fig += 1
            else:
                _add_runs(doc.add_paragraph(), stripped)
            i += 1
            continue

        # Madde
        if re.match(r"^[-*]\s+", stripped):
            _add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            _add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Bos satir
        if not stripped:
            i += 1
            continue

        # Normal paragraf
        _add_runs(doc.add_paragraph(), stripped)
        i += 1

    doc.save(str(OUT))
    print(f"OK -> {OUT}")
    print(f"   {n_tab} tablo, {n_fig} figur gomuldu.")
    print("   UYARI TASLAK: Aptos/marj/<=10-sayfa Word'de dogrulanmali; tablo/figur tasmasi elle ayarlanir.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
